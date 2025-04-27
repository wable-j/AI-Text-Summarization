"""
Utility functions for text extraction from different sources
"""
import requests
from bs4 import BeautifulSoup
import fitz  # PyMuPDF
import io

def extract_text_from_url(url):
    """
    Extract text content from a webpage
    
    Args:
        url (str): The URL to fetch and extract text from
        
    Returns:
        str: Extracted text content
    """
    try:
        response = requests.get(url)
        response.raise_for_status()  # Raise exception for 4XX/5XX responses
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Remove script and style elements
        for script_or_style in soup(["script", "style"]):
            script_or_style.decompose()
        
        # Get text from paragraphs, headings, and other relevant tags
        text_elements = soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'article'])
        
        # Extract text and join with spaces
        text = ' '.join([elem.get_text(strip=True) for elem in text_elements])
        
        # Clean up whitespace
        text = ' '.join(text.split())
        
        return text
    except Exception as e:
        raise Exception(f"Failed to extract text from URL: {str(e)}")

def extract_text_from_pdf(file_content):
    """
    Extract text content from a PDF file
    
    Args:
        file_content (bytes): The PDF file content
        
    Returns:
        str: Extracted text content
    """
    try:
        pdf_file = fitz.open(stream=io.BytesIO(file_content), filetype="pdf")
        text = ""
        
        for page_num in range(len(pdf_file)):
            page = pdf_file[page_num]
            text += page.get_text()
        
        # Clean up whitespace
        text = ' '.join(text.split())
        
        return text
    except Exception as e:
        raise Exception(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(file_content):
    """
    Extract text content from a DOCX file
    
    Args:
        file_content (bytes): The DOCX file content
        
    Returns:
        str: Extracted text content
    """
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_content))
        
        # Extract text from paragraphs
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + ' '
        
        # Clean up whitespace
        text = ' '.join(text.split())
        
        return text
    except Exception as e:
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")

def extract_text_from_html(html_content):
    """
    Extract text content from HTML content
    
    Args:
        html_content (str): The HTML content
        
    Returns:
        str: Extracted text content
    """
    try:
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script_or_style in soup(["script", "style"]):
            script_or_style.decompose()
        
        # Get text from the HTML
        text = soup.get_text()
        
        # Clean up whitespace
        text = ' '.join(text.split())
        
        return text
    except Exception as e:
        raise Exception(f"Failed to extract text from HTML: {str(e)}")